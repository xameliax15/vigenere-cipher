from docx import Document
from docx.shared import Inches

# Data for the table
rows = [
    ["1", "Buat CMS sederhana dengan UI Admin LTE", "Code Awal CMS", "Code awal hasil generate AI"],
    ["2", "Perbaiki koneksi database dan setup struktur tabel", "Fix Database Connection & Setup Tables", "Memperbaiki koneksi MySQL dan membuat struktur tabel"],
    ["3", "Implementasi fitur login, register, dan autentikasi user", "Add User Authentication", "Menambah fitur login, register, dan autentikasi user"],
    ["4", "Tambahkan fitur CRUD untuk post, page, dan media", "Add CRUD for Posts, Pages, and Media", "Menambah fitur CRUD untuk post, page, dan media"],
    ["5", "Ubah tampilan login dan dashboard agar lebih modern dan biru", "Update Login & Dashboard UI", "Membuat tampilan login dan dashboard lebih modern"],
    ["6", "Ubah tampilan posts, pages, media, dan profile agar konsisten dan modern", "Update UI for Posts, Pages, Media, and Profile", "Menyamakan tampilan posts, pages, media, dan profile"]
]

# Create a new Word document
doc = Document()
doc.add_paragraph("Riwayat commit ke repo :")

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Perintah Ke AI'
hdr_cells[2].text = 'Judul Commit'
hdr_cells[3].text = 'Deskripsi Commit'

for row in rows:
    row_cells = table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = val

# Save the document
doc.save('commit_history.docx') 