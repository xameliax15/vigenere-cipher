from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

# Data commit
rows = [
    ["1", "Buat CMS sederhana dengan\nUI Admin  LTE", "Code Awal CMS", "Code awal hasil generate AI"],
    ["2", "Perbaiki koneksi database dan setup struktur tabel", "Fix Database Connection & Setup Tables", "Memperbaiki koneksi MySQL dan membuat struktur tabel"],
    ["3", "Implementasi fitur login, register, dan autentikasi user", "Add User Authentication", "Menambah fitur login, register, dan autentikasi user"],
    ["4", "Tambahkan fitur CRUD untuk post, page, dan media", "Add CRUD for Posts, Pages, and Media", "Menambah fitur CRUD untuk post, page, dan media"],
    ["5", "Ubah tampilan login dan dashboard agar lebih modern dan biru", "Update Login & Dashboard UI", "Membuat tampilan login dan dashboard lebih modern"],
    ["6", "Ubah tampilan posts, pages, media, dan profile agar konsisten dan modern", "Update UI for Posts, Pages, Media, and Profile", "Menyamakan tampilan posts, pages, media, dan profile"]
]

doc = Document()

# Header identitas dan UTS
header_table = doc.add_table(rows=4, cols=4)
header_table.style = 'Table Grid'
header_table.autofit = False
header_table.columns[0].width = Inches(1.2)
header_table.columns[1].width = Inches(2.2)
header_table.columns[2].width = Inches(1.2)
header_table.columns[3].width = Inches(2.2)

header_table.cell(0,0).text = 'Nama'
header_table.cell(0,1).text = ':'
header_table.cell(0,2).merge(header_table.cell(0,3)).text = 'UJIAN TENGAH SEMESTER\nFAKULTAS TEKNOLOGI INFORMASI'
header_table.cell(1,0).text = 'NIM'
header_table.cell(1,1).text = ':'
header_table.cell(1,2).merge(header_table.cell(1,3)).text = ''
header_table.cell(2,0).text = 'Kelas'
header_table.cell(2,1).text = ':'
header_table.cell(2,2).merge(header_table.cell(2,3)).text = ''
header_table.cell(3,0).text = 'Program Studi'
header_table.cell(3,1).text = ':'
header_table.cell(3,2).merge(header_table.cell(3,3)).text = ''

# Soal
p = doc.add_paragraph()
p.add_run('Soal :').bold = True
soal_table = doc.add_table(rows=1, cols=1)
soal_table.style = 'Table Grid'
soal_table.cell(0,0).text = '\nTuliskan Kembali Soal\n'

# Link ke repo
p = doc.add_paragraph()
p.add_run('Link ke repo Github :').bold = True
link_table = doc.add_table(rows=1, cols=1)
link_table.style = 'Table Grid'
link_table.cell(0,0).text = '\nTuliskan Link Repo Github Anda\n'

# Riwayat commit
p = doc.add_paragraph()
p.add_run('Riwayat commit ke repo :').bold = True

commit_table = doc.add_table(rows=1, cols=4)
commit_table.style = 'Table Grid'
commit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = commit_table.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Perintah Ke AI'
hdr_cells[2].text = 'Judul Commit'
hdr_cells[3].text = 'Deskripsi Commit'

for row in rows:
    row_cells = commit_table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = val

# Tambahkan baris kosong jika kurang dari 6
for _ in range(len(rows), 6):
    row_cells = commit_table.add_row().cells
    for i in range(4):
        row_cells[i].text = ''

# Save
doc.save('UTS_Format_Commit.docx') 